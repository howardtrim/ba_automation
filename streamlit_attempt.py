import streamlit as st
import pandas as pd
import docx
from docx import Document

st.write('Sweeties Words')

base = pd.read_csv('concat_base.csv')
base.columns = ['level', 'text']

job_type = pd.read_csv('Jobs.csv', header = None)
job_type.columns = ['jobs']
types = list(job_type.jobs)

manager = list(base.level.unique())

project_name = st.text_input("Enter project name")
project_number = st.text_input("Enter project number")
fieldwork_start = st.text_input("Enter fieldwork start date")
project_deadline = st.text_input("Enter project deadline")
data_source = st.text_input("Data to be sourced by")
data_deadline = st.text_input("Data check deadline")

sample_size = st.text_input("Sample size prior to the full data sourcing")
no_contacts = st.text_input("Quantity of companies to search")
country_list = st.text_input("List of countries to search")
minimum_contacts = st.text_input("Minimum number of contacts to search")
maximum_contacts = st.text_input("Maximum number of contacts to search")

minmax_contacts = 'Minimum: ' + minimum_contacts + ', Maximum: ' + maximum_contacts

level = st.selectbox('Select Managerial Level', manager)
department = st.multiselect('Select Department',types)

level2 = st.selectbox('Select Managerial for another Level', manager)
department2 = st.multiselect('Select other Departments',types)

concat = base[base.level == level]
joiner  = pd.DataFrame(concat['text']).merge(pd.DataFrame(department), how ='cross')
joiner['keyword'] = joiner.agg(' '.join, axis=1)

concat2 = base[base.level == level2]
joiner2  = pd.DataFrame(concat2['text']).merge(pd.DataFrame(department2), how ='cross')
joiner2['keyword'] = joiner2.agg(' '.join, axis=1)

keywords = list(joiner['keyword'])
keywords = ", ".join(keywords)

keywords2 = list(joiner2['keyword'])
keywords2 = ", ".join(keywords2)

st.write('these are the keywords')
st.write(keywords)

st.write('these are the other keywords')
st.write(keywords2)

all_level = level + ', ' + level2
all_department = department + department2
all_department = ", ".join(all_department)

additional_notes = st.text_input("Additional start-up notes")
file_name = st.text_input("Name the file")

if file_name != '':

    document = Document("Sourcing_Data_Request_Form.docx")

    document.tables[0].cell(0,0).add_paragraph(project_name)
    document.tables[0].cell(0,1).add_paragraph(project_number)
    document.tables[0].cell(0,2).add_paragraph(fieldwork_start)
    document.tables[0].cell(0,3).add_paragraph(project_deadline)
    document.tables[1].cell(0,0).add_paragraph(data_source)
    document.tables[1].cell(0,1).add_paragraph(data_deadline)
    document.tables[2].cell(0,1).add_paragraph(additional_notes)

    document.paragraphs[6].add_run(text = sample_size )
    document.paragraphs[7].add_run(text = no_contacts)
    document.paragraphs[8].add_run(text = country_list)
    document.paragraphs[9].add_run(text = minmax_contacts)

    document.paragraphs[20].add_run(text = all_level)
    document.paragraphs[23].add_run(text = all_department)

    document.paragraphs[27].add_run(text = keywords)
    document.paragraphs[28].add_run(text = keywords2)
    document.save(file_name + '.docx')
